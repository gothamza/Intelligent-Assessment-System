"""
Exam Generation Page
Allows users to generate exams with percentage distribution for each course
"""
import streamlit as st
import requests
import os
import json
from datetime import datetime
import io
from data.moroccan_curriculum import get_courses_for_grade_subject
from utils.translations import get_translation, get_translations

BACKEND_URL = os.getenv("BACKEND_URL", "http://backend:8000")

def generate_exam_template_pdf(exam, include_answers=False):
    """Generate PDF from exam template (without student answers)"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, 
                               rightMargin=72, leftMargin=72, 
                               topMargin=72, bottomMargin=72)
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Heading1'],
            fontSize=20,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#1f4788'),
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'HeadingStyle',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.HexColor('#2c3e50'),
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'NormalStyle',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            alignment=TA_LEFT,
            leading=14
        )
        
        question_style = ParagraphStyle(
            'QuestionStyle',
            parent=normal_style,
            fontSize=12,
            fontName='Helvetica-Bold',
            textColor=colors.HexColor('#34495e')
        )
        
        answer_style = ParagraphStyle(
            'AnswerStyle',
            parent=normal_style,
            fontSize=10,
            leftIndent=20,
            textColor=colors.HexColor('#27ae60'),
            backColor=colors.HexColor('#ecf0f1')
        )
        
        story = []
        
        # Title
        exam_title = exam.get('title', 'Examen')
        story.append(Paragraph(exam_title, title_style))
        story.append(Spacer(1, 10))
        
        # Exam info
        info_data = [
            ['Matière:', exam.get('subject', 'N/A')],
            ['Niveau:', exam.get('grade', 'N/A')],
            ['Difficulté:', exam.get('difficulty', 'N/A')],
            ['Date:', datetime.now().strftime('%d/%m/%Y')]
        ]
        info_table = Table(info_data, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#ecf0f1')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 12),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 20))
        
        # Instructions
        story.append(Paragraph("Instructions:", heading_style))
        instructions_text = exam.get('instructions', 'Répondez à toutes les questions. Justifiez vos réponses clairement.')
        story.append(Paragraph(instructions_text, normal_style))
        story.append(Spacer(1, 20))
        
        # Questions
        questions = exam.get('questions', [])
        
        for i, question_data in enumerate(questions):
            story.append(Paragraph(f"Question {i+1}", heading_style))
            
            # Course and difficulty
            course = question_data.get('course', 'N/A')
            difficulty = question_data.get('difficulty', 'N/A')
            story.append(Paragraph(f"<b>Cours:</b> {course} | <b>Difficulté:</b> {difficulty}", normal_style))
            story.append(Spacer(1, 5))
            
            # Question text
            question_text = question_data.get('question', '').replace('\n', '<br/>')
            story.append(Paragraph(f"<b>Question:</b> {question_text}", question_style))
            story.append(Spacer(1, 10))
            
            # Answer space (blank lines for student to write)
            story.append(Paragraph("Réponse:", normal_style))
            story.append(Spacer(1, 5))
            story.append(Paragraph("_" * 80, normal_style))
            story.append(Spacer(1, 5))
            story.append(Paragraph("_" * 80, normal_style))
            story.append(Spacer(1, 5))
            story.append(Spacer(1, 10))
            
            # Answer (if included)
            if include_answers and question_data.get('answer'):
                answer_text = str(question_data.get('answer', '')).replace('\n', '<br/>')
                story.append(Paragraph(f"<b>Réponse:</b> {answer_text}", answer_style))
                story.append(Spacer(1, 10))
            
            story.append(Spacer(1, 15))
            
            # Page break every 3 questions
            if (i + 1) % 3 == 0 and i < len(questions) - 1:
                story.append(PageBreak())
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        st.error(f"Erreur lors de la génération du PDF: {str(e)}")
        return None

def generate_exam_template_docx(exam, include_answers=False):
    """Generate DOCX from exam template (without student answers)"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Title
        title = doc.add_heading(exam.get('title', 'Examen'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(20)
        title_run.font.color.rgb = RGBColor(31, 71, 136)
        
        # Exam info
        doc.add_paragraph()
        info_para = doc.add_paragraph()
        info_para.add_run('Matière: ').bold = True
        info_para.add_run(exam.get('subject', 'N/A'))
        info_para.add_run(' | Niveau: ').bold = True
        info_para.add_run(exam.get('grade', 'N/A'))
        info_para.add_run(' | Difficulté: ').bold = True
        info_para.add_run(exam.get('difficulty', 'N/A'))
        info_para.add_run(' | Date: ').bold = True
        info_para.add_run(datetime.now().strftime('%d/%m/%Y'))
        
        doc.add_paragraph()
        
        # Instructions
        instructions_heading = doc.add_heading('Instructions', level=2)
        instructions_heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
        instructions_text = exam.get('instructions', 'Répondez à toutes les questions. Justifiez vos réponses clairement.')
        doc.add_paragraph(instructions_text)
        
        doc.add_paragraph()
        doc.add_paragraph('─' * 50)
        doc.add_paragraph()
        
        # Questions
        questions = exam.get('questions', [])
        
        for i, question_data in enumerate(questions):
            # Question number
            q_heading = doc.add_heading(f'Question {i+1}', level=2)
            q_heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            
            # Course and difficulty
            course_para = doc.add_paragraph()
            course_para.add_run('Cours: ').bold = True
            course_para.add_run(question_data.get('course', 'N/A'))
            course_para.add_run(' | Difficulté: ').bold = True
            course_para.add_run(question_data.get('difficulty', 'N/A'))
            
            # Question text
            doc.add_paragraph()
            q_para = doc.add_paragraph()
            q_para.add_run('Question: ').bold = True
            q_para.add_run(question_data.get('question', ''))
            
            # Answer space (blank lines)
            doc.add_paragraph()
            ans_para = doc.add_paragraph()
            ans_para.add_run('Réponse: ').bold = True
            doc.add_paragraph('_' * 80)
            doc.add_paragraph('_' * 80)
            doc.add_paragraph()
            
            # Answer (if included)
            if include_answers and question_data.get('answer'):
                answer_para = doc.add_paragraph()
                answer_para.add_run('Réponse: ').bold = True
                answer_run = answer_para.add_run(str(question_data.get('answer', '')))
                answer_run.font.color.rgb = RGBColor(39, 174, 96)  # Green
                answer_para.paragraph_format.left_indent = Inches(0.5)
                answer_para.paragraph_format.space_before = Pt(6)
                # Add background color
                for run in answer_para.runs:
                    run.font.highlight_color = 7  # Light gray highlight
            
            doc.add_paragraph()
            doc.add_paragraph('─' * 50)
            doc.add_paragraph()
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        # Error message is not user-facing, can stay in French for debugging
        st.error(f"Erreur lors de la génération du DOCX: {str(e)}")
        return None

def generate_exam_with_percentages():
    """Main function for exam generation page"""
    
    # Get current language for translations
    language = st.session_state.get("selected_language", "Français")
    t = get_translations(language, "exam_generation")
    
    title_text = t.get('title', "Création d'Examen")
    st.title(f"📝 {title_text}")
    st.markdown("---")
    st.markdown(t.get('subtitle', "Générez un examen personnalisé avec une répartition par cours."))
    
    headers = {"Authorization": f"Bearer {st.session_state.get('token')}"}
    
    # Check authentication
    if not st.session_state.get("token"):
        error_msg = get_translation(language, "common", "error", "Erreur")
        st.error(f"❌ {error_msg}")
        st.info(get_translation(language, "common", "info", "Information"))
        return
    
    # Initialize session state for exam
    if "generated_exam" not in st.session_state:
        st.session_state.generated_exam = None
    if "exam_title" not in st.session_state:
        st.session_state.exam_title = ""
    
    # Get configuration from sidebar (subject, grade)
    # Language will be retrieved dynamically when generating the exam
    subject = st.session_state.get("selected_subject", "Mathématiques")
    grade = st.session_state.get("selected_grade", "7ème année")
    
    # Course selection section
    st.subheader(f"📚 {t.get('course_selection', 'Sélection des Cours')}")
    
    # Get available courses for selected grade and subject
    available_courses = get_courses_for_grade_subject(grade, subject)
    
    if not available_courses:
        st.warning(f"⚠️ {t.get('no_courses', 'Aucun cours disponible pour')} {subject} - {grade}")
        st.info(get_translation(language, "common", "info", "Information"))
        return
    
    # Allow user to select multiple courses
    selected_courses = st.multiselect(
        t.get('select_courses', "Choisissez les cours à inclure dans l'examen:"),
        options=available_courses,
        default=available_courses[:min(3, len(available_courses))] if available_courses else [],
        help=t.get('select_courses', "Sélectionnez un ou plusieurs cours pour l'examen")
    )
    
    if not selected_courses:
        st.info(f"💡 {get_translation(language, 'common', 'info', 'Information')}")
        return
    
    # Percentage distribution section
    st.markdown("---")
    course_dist_label = get_translation(language, "exam_generation", "course_percentage", "Répartition des Questions par Cours")
    st.subheader(f"📊 {course_dist_label}")
    
    num_questions = st.number_input(
        get_translation(language, "exam_generation", "num_questions", "Nombre total de questions:"),
        min_value=1,
        max_value=30,
        value=5,
        step=1,
        help=get_translation(language, "exam_generation", "num_questions", "Nombre total de questions dans l'examen")
    )
    
    # Mode selection: Auto or Manual
    mode_auto = t.get('percentage_mode_auto', "Automatique (égale)")
    mode_manual = t.get('percentage_mode_manual', "Manuelle (personnalisée)")
    mode = st.radio(
        get_translation(language, "exam_generation", "course_percentage", "Mode de répartition:"),
        [f"🎯 {mode_auto}", f"✏️ {mode_manual}"],
        horizontal=True,
        help=get_translation(language, "exam_generation", "course_percentage", "Choisissez comment répartir les questions entre les cours")
    )
    
    course_percentages = {}
    
    if f"🎯 {mode_auto}" in mode:
        # Equal distribution
        equal_percentage = 100 // len(selected_courses)
        remainder = 100 % len(selected_courses)
        
        for i, course in enumerate(selected_courses):
            percentage = equal_percentage + (1 if i < remainder else 0)
            course_percentages[course] = percentage
        
        # Display auto distribution
        auto_dist_label = get_translation(language, "exam_generation", "percentage_mode_auto", "Répartition automatique")
        st.success(f"✅ {auto_dist_label}:")
        question_label = get_translation(language, "exam_generation", "num_questions", "question(s)")
        for course, percentage in course_percentages.items():
            num_q = max(1, int((percentage / 100) * num_questions))
            st.write(f"  • **{course}**: {percentage}% ({num_q} {question_label})")
    else:
        # Manual distribution
        manual_dist_label = get_translation(language, "exam_generation", "course_percentage", "Définissez le pourcentage pour chaque cours")
        st.markdown(f"**{manual_dist_label}:**")
        
        auto_normalize = st.checkbox(
            t.get('auto_normalize', "Auto-normaliser à 100%"),
            value=True,
            help=t.get('auto_normalize', "Ajuste automatiquement les pourcentages pour totaliser 100%")
        )
        
        total_percentage = 0
        
        for course in selected_courses:
            # Calculate initial value for even distribution
            initial_value = 100 // len(selected_courses)
            
            # Get previous value if exists
            key = f"exam_percentage_{course}"
            if key in st.session_state:
                initial_value = st.session_state[key]
            
            percentage = st.slider(
                f"{course} (%)",
                min_value=0,
                max_value=100,
                value=initial_value,
                step=1,
                key=key,
                help=f"Pourcentage de questions pour {course}"
            )
            
            course_percentages[course] = percentage
            total_percentage += percentage
        
        # Normalize if enabled
        if auto_normalize and total_percentage != 100 and total_percentage > 0:
            # Scale all percentages to total 100
            scale_factor = 100 / total_percentage
            for course in course_percentages:
                course_percentages[course] = round(course_percentages[course] * scale_factor, 1)
            total_percentage = sum(course_percentages.values())
            # Rounding adjustment
            diff = 100 - total_percentage
            if diff != 0 and selected_courses:
                course_percentages[selected_courses[0]] += diff
        
        # Display total
        total_percentage = sum(course_percentages.values())
        total_label = t.get('total_percentage', "Total")
        must_be_100 = t.get('must_be_100', "doit être 100%")
        perfect = t.get('perfect', "Parfait!")
        if total_percentage != 100:
            st.warning(f"⚠️ {total_label}: {total_percentage}% ({must_be_100})")
        else:
            st.success(f"✅ {total_label}: {total_percentage}% - {perfect}")
        
        # Display preview
        preview_label = t.get('preview_distribution', "Aperçu de la répartition:")
        st.markdown(f"**{preview_label}**")
        for course, percentage in course_percentages.items():
            num_q = max(1, round((percentage / 100) * num_questions))
            st.write(f"  • **{course}**: {percentage}% ({num_q} {get_translation(language, 'exam_generation', 'num_questions', 'question(s)')})")
    
    # Additional settings
    st.markdown("---")
    settings_label = t.get('exam_settings', "Paramètres de l'Examen")
    st.subheader(f"⚙️ {settings_label}")
    
    difficulty_options = [
        get_translation(language, "common", "difficulty_easy", "Facile"),
        get_translation(language, "common", "difficulty_intermediate", "Intermédiaire"),
        get_translation(language, "common", "difficulty_hard", "Difficile")
    ]
    difficulty = st.selectbox(
        t.get('difficulty', "Difficulté:"),
        difficulty_options,
        index=1,
        help=get_translation(language, "exam_generation", "difficulty", "Niveau de difficulté des questions")
    )
    
    exam_title = st.text_input(
        t.get('exam_title', "Titre de l'examen:"),
        value=f"Examen de {subject} - {grade} - {difficulty}",
        help=t.get('exam_title', "Titre personnalisé pour l'examen")
    )
    
    include_answers = st.checkbox(
        t.get('include_answers', "Inclure les réponses"),
        value=True,
        help=t.get('include_answers', "Générer les réponses avec les questions")
    )
    
    # Generate exam button
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        generate_btn_label = t.get('generate_exam', "Générer l'Examen")
        if st.button(f"🚀 {generate_btn_label}", type="primary", use_container_width=True):
            if sum(course_percentages.values()) != 100 and f"✏️ {mode_manual}" in mode:
                error_msg = t.get('error_percentage', "La somme des pourcentages doit être égale à 100%")
                st.error(f"❌ {error_msg}")
            else:
                # Get language dynamically from session state
                current_language = st.session_state.get("selected_language", "Français")
                generate_exam(
                    headers,
                    subject,
                    grade,
                    selected_courses,
                    course_percentages,
                    num_questions,
                    difficulty,
                    exam_title,
                    include_answers,
                    current_language
                )
    
    # Display generated exam
    if st.session_state.generated_exam:
        st.markdown("---")
        generated_label = t.get('generated_exam', "Examen Généré")
        st.subheader(f"📋 {generated_label}")
        
        exam = st.session_state.generated_exam
        
        # Exam header
        st.markdown(f"### {exam.get('title', exam_title)}")
        subject_label = t.get('subject', "Matière:")
        grade_label = t.get('grade', "Niveau:")
        difficulty_label = t.get('difficulty', "Difficulté:")
        st.markdown(f"**{subject_label}** {subject} | **{grade_label}** {grade} | **{difficulty_label}** {difficulty}")
        num_q_label = t.get('num_questions', "Nombre de questions:")
        st.markdown(f"**{num_q_label}** {len(exam.get('questions', []))}")
        st.markdown("---")
        
        # Display questions
        question_label = t.get('question', "Question:")
        course_label = t.get('course', "Cours:")
        answer_label = t.get('answer', "Réponse:")
        for i, question_data in enumerate(exam.get('questions', []), 1):
            with st.expander(f"{question_label} {i}: {question_data.get('course', 'N/A')}"):
                st.markdown(f"**{course_label}** {question_data.get('course', 'N/A')}")
                st.markdown(f"**{question_label}**\n{question_data.get('question', 'N/A')}")
                
                if include_answers and question_data.get('answer'):
                    st.markdown(f"**{answer_label}**\n{question_data.get('answer', 'N/A')}")
        
        # Download options
        st.markdown("---")
        download_label = t.get('download_exam', "Télécharger l'Examen")
        st.subheader(f"💾 {download_label}")
        
        col_d1, col_d2, col_d3, col_d4 = st.columns(4)
        
        with col_d1:
            # JSON download
            exam_json = json.dumps(exam, ensure_ascii=False, indent=2)
            st.download_button(
                label="📄 JSON",
                data=exam_json,
                file_name=f"examen_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json",
                use_container_width=True,
                help="Télécharger l'examen au format JSON"
            )
        
        with col_d2:
            # PDF download
            pdf_data = generate_exam_template_pdf(exam, include_answers=include_answers)
            if pdf_data:
                st.download_button(
                    label="📕 PDF",
                    data=pdf_data,
                    file_name=f"examen_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    help="Télécharger l'examen au format PDF"
                )
            else:
                st.button("📕 PDF", disabled=True, use_container_width=True, help="Erreur lors de la génération")
        
        with col_d3:
            # Word download
            docx_data = generate_exam_template_docx(exam, include_answers=include_answers)
            if docx_data:
                st.download_button(
                    label="📘 Word",
                    data=docx_data,
                    file_name=f"examen_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    help="Télécharger l'examen au format Word (DOCX)"
                )
            else:
                st.button("📘 Word", disabled=True, use_container_width=True, help="Erreur lors de la génération")
        
        with col_d4:
            # Save to session state for exam taking page
            take_exam_label = t.get('take_exam', "Passer l'Examen")
            if st.button(f"📝 {take_exam_label}", use_container_width=True):
                st.session_state.exam_to_take = exam
                exam_saved_msg = t.get('exam_saved', "Examen sauvegardé! Allez à la page 'Passer un Examen' pour commencer.")
                st.success(f"✅ {exam_saved_msg}")
                nav_msg = t.get('use_navigation', "Utilisez la navigation latérale pour accéder à la page 'Passer un Examen'.")
                st.info(f"💡 {nav_msg}")


def generate_exam(headers, subject, grade, courses, course_percentages, num_questions, difficulty, exam_title, include_answers, language):
    """Generate exam using the backend"""
    
    # Get translations for this function
    t_gen = get_translations(language, "exam_generation")
    
    generating_msg = t_gen.get('generating', "Génération de l'examen en cours...")
    with st.spinner(f"🤔 {generating_msg}"):
        try:
            # Calculate number of questions per course
            course_counts = {}
            for course, percentage in course_percentages.items():
                count = max(1, round((percentage / 100) * num_questions))
                course_counts[course] = count
            
            # Adjust to match total
            total = sum(course_counts.values())
            if total != num_questions:
                diff = num_questions - total
                # Add/subtract from largest course
                largest_course = max(course_counts, key=course_counts.get)
                course_counts[largest_course] += diff
            
            # Build prompt for exam generation
            # Translate course names and labels based on language
            course_list_parts = []
            if language == "English":
                course_translations = {
                    "Nombres entiers et décimaux": "Integers and Decimals",
                    "Opérations de base": "Basic Operations",
                    "Géométrie plane": "Plane Geometry",
                    "Fractions": "Fractions",
                    "Pourcentages": "Percentages",
                    "Proportions": "Proportions"
                }
                for course, count in course_counts.items():
                    translated_course = course_translations.get(course, course)
                    course_list_parts.append(f"{translated_course} ({count} question(s))")
            elif language == "العربية":
                course_translations = {
                    "Nombres entiers et décimaux": "الأعداد الصحيحة والعشرية",
                    "Opérations de base": "العمليات الأساسية",
                    "Géométrie plane": "الهندسة المستوية",
                    "Fractions": "الكسور",
                    "Pourcentages": "النسب المئوية",
                    "Proportions": "النسب"
                }
                for course, count in course_counts.items():
                    translated_course = course_translations.get(course, course)
                    course_list_parts.append(f"{translated_course} ({count} سؤال(ات))")
            else:
                # French (default)
                for course, count in course_counts.items():
                    course_list_parts.append(f"{course} ({count} question(s))")
            
            course_list_text = ", ".join(course_list_parts)
            
            # Fix f-string issue: extract answer field logic
            answer_field_fr = '"answer": "réponse_détaillée",' if include_answers else '"answer": null,'
            answer_field_en = '"answer": "detailed_answer",' if include_answers else '"answer": null,'
            answer_field_ar = '"answer": "إجابة_تفصيلية",' if include_answers else '"answer": null,'
            
            answer_instruction_fr = "La réponse détaillée avec explications" if include_answers else "Structure pour la réponse (sans donner la réponse complète)"
            answer_instruction_en = "The detailed answer with explanations" if include_answers else "Structure for the answer (without providing the complete answer)"
            answer_instruction_ar = "الإجابة التفصيلية مع الشروحات" if include_answers else "هيكل للإجابة (بدون إعطاء الإجابة الكاملة)"
            
            # Map difficulty to English if needed
            difficulty_en = difficulty
            if language == "English":
                difficulty_map = {
                    "Facile": "Easy",
                    "Intermédiaire": "Intermediate",
                    "Difficile": "Hard",
                    "Easy": "Easy",
                    "Intermediate": "Intermediate",
                    "Hard": "Hard"
                }
                difficulty_en = difficulty_map.get(difficulty, difficulty)
            
            # Translate exam title template based on language
            if language == "English":
                title_template = f"Mathematics Exam - 7th Grade - {difficulty_en}"
                subject_translated = "Mathematics"
                grade_translated = "7th Grade"
            elif language == "العربية":
                title_template = f"امتحان الرياضيات - الصف السابع - {difficulty}"
                subject_translated = "الرياضيات"
                grade_translated = "الصف السابع"
            else:
                title_template = exam_title
                subject_translated = subject
                grade_translated = grade
            
            prompts = {
                "Français": f"""Génère-moi un examen de {subject} pour le niveau {grade} avec les spécifications suivantes:
- Difficulté: {difficulty}
- Nombre total de questions: {num_questions}
- Répartition par cours: {course_list_text}

Pour chaque question, fournis:
1. Le cours correspondant
2. La question complète et claire
3. {answer_instruction_fr}

Format de sortie JSON:
{{
  "title": "{exam_title}",
  "subject": "{subject}",
  "grade": "{grade}",
  "difficulty": "{difficulty}",
  "questions": [
    {{
      "course": "nom_du_cours",
      "question": "texte_de_la_question",
      {answer_field_fr}
      "difficulty": "{difficulty}"
    }}
  ]
}}

Assure-toi que la répartition des questions respecte exactement les nombres spécifiés pour chaque cours.""",
                
                "English": f"""Generate an exam in {subject_translated} for {grade_translated} level with the following specifications:
- Difficulty: {difficulty_en}
- Total number of questions: {num_questions}
- Distribution by course: {course_list_text}

IMPORTANT: You MUST generate ALL content (title, questions, answers, course names) in English ONLY. Do not use French words.

For each question, provide:
1. The corresponding course (translate course names to English)
2. The complete and clear question (in English)
3. {answer_instruction_en}

Output format JSON:
{{
  "title": "{title_template}",
  "subject": "{subject_translated}",
  "grade": "{grade_translated}",
  "difficulty": "{difficulty_en}",
  "questions": [
    {{
      "course": "course_name_in_english",
      "question": "question_text_in_english",
      {answer_field_en}
      "difficulty": "{difficulty_en}"
    }}
  ]
}}

Make sure the question distribution exactly matches the specified numbers for each course. Remember: ALL content must be in English, including course names.""",
                
                "العربية": f"""أنشئ لي امتحانًا في {subject} لمستوى {grade} بالمواصفات التالية:
- الصعوبة: {difficulty}
- العدد الإجمالي للأسئلة: {num_questions}
- التوزيع حسب المادة: {course_list_text}

لكل سؤال، قدم:
1. المادة المقابلة
2. السؤال الكامل والواضح
3. {answer_instruction_ar}

تنسيق الإخراج JSON:
{{
  "title": "{exam_title}",
  "subject": "{subject}",
  "grade": "{grade}",
  "difficulty": "{difficulty}",
  "questions": [
    {{
      "course": "اسم_المادة",
      "question": "نص_السؤال",
      {answer_field_ar}
      "difficulty": "{difficulty}"
    }}
  ]
}}

تأكد من أن توزيع الأسئلة يطابق تمامًا الأعداد المحددة لكل مادة."""
            }
            
            prompt = prompts.get(language, prompts["Français"])
            
            # Create a temporary chat for exam generation
            chat_resp = requests.post(
                f"{BACKEND_URL}/chats/create",
                json={"chat_title": f"Examen {exam_title}"},
                headers=headers
            )
            
            if chat_resp.status_code != 200:
                error_msg = get_translation(language, "common", "error", "Erreur")
                st.error(f"❌ {error_msg}")
                return
            
            chat_id = chat_resp.json().get("chat_id")
            
            # Prepare messages for backend
            messages = [{"role": "user", "content": prompt}]
            
            # Call RAG backend to generate exam
            body = {
                "doc_ids": [],
                "chat_id": chat_id,
                "messages": messages,
                "use_web_search": False,
                "model_name": st.session_state.get("selected_model", "Groq/llama-3.1-8b-instant"),
                "subject": subject,
                "grade": grade,
                "course": ", ".join(courses),
                "custom_instructions": "",
                "language": language
            }
            
            res = requests.post(
                f"{BACKEND_URL}/rag_langgraph2",
                json=body,
                headers=headers,
                timeout=300
            )
            
            if res.status_code == 200:
                data = res.json()
                ai_response = data.get("answer", "")
                
                # Try to extract JSON from response with robust parsing
                try:
                    import re
                    
                    # Function to extract and fix JSON with LaTeX support
                    def extract_and_fix_json(text):
                        """Extract JSON from text and fix LaTeX backslashes"""
                        # First, try to find JSON in markdown code blocks
                        json_code_block = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', text, re.DOTALL)
                        if json_code_block:
                            json_str = json_code_block.group(1)
                        else:
                            # Try to find JSON object in the response (looking for opening and closing braces)
                            # More robust: find first { and matching }
                            brace_count = 0
                            start_idx = text.find('{')
                            if start_idx == -1:
                                return None
                            
                            end_idx = start_idx
                            for i in range(start_idx, len(text)):
                                if text[i] == '{':
                                    brace_count += 1
                                elif text[i] == '}':
                                    brace_count -= 1
                                    if brace_count == 0:
                                        end_idx = i + 1
                                        break
                            
                            if brace_count == 0:
                                json_str = text[start_idx:end_idx]
                            else:
                                # Fallback: use regex
                                json_match = re.search(r'\{.*\}', text, re.DOTALL)
                                if json_match:
                                    json_str = json_match.group(0)
                                else:
                                    return None
                        
                        # Fix LaTeX backslashes in string values
                        # Strategy: First try to parse, if it fails with escape error, fix and retry
                        try:
                            # Try parsing directly first (might work if no LaTeX)
                            return json.loads(json_str)
                        except json.JSONDecodeError as e:
                            if "Invalid \\escape" in str(e) or "Invalid escape" in str(e):
                                # We have an escape error - fix LaTeX backslashes
                                # Strategy: Protect valid JSON escapes, escape LaTeX commands, restore escapes
                                
                                # Fix LaTeX backslashes by escaping them BEFORE protecting valid escapes
                                # The issue: \begin contains \b (valid escape), but \b should NOT be interpreted as backspace
                                # Solution: Escape LaTeX commands (2+ letters) FIRST, then protect valid single-char escapes
                                
                                # Step 1: Store unicode escapes (\uXXXX) first (they're unambiguous)
                                unicode_storage = {}
                                counter = [0]
                                
                                def store_unicode(match):
                                    marker = f'__UNI_{counter[0]}__'
                                    unicode_storage[marker] = match.group(0)
                                    counter[0] += 1
                                    return marker
                                
                                temp_json = json_str
                                temp_json = re.sub(r'\\u[0-9A-Fa-f]{4}', store_unicode, temp_json)
                                
                                # Step 2: Escape LaTeX commands (2+ letters after \) BEFORE protecting valid escapes
                                # This way, \begin becomes \\begin before we protect \b
                                temp_json = re.sub(r'\\([a-zA-Z]{2,})', r'\\\\\1', temp_json)
                                
                                # Step 3: Now protect valid single-char JSON escapes
                                valid_escapes = {
                                    '\\n': '__ESC_N__',
                                    '\\t': '__ESC_T__',
                                    '\\r': '__ESC_R__',
                                    '\\"': '__ESC_Q__',
                                    '\\\\': '__ESC_BS__',
                                    '\\/': '__ESC_SL__',
                                    '\\b': '__ESC_B__',
                                    '\\f': '__ESC_F__',
                                }
                                
                                for escape, marker in valid_escapes.items():
                                    temp_json = temp_json.replace(escape, marker)
                                
                                # Step 4: Restore valid escapes
                                final_json = temp_json
                                for escape, marker in valid_escapes.items():
                                    final_json = final_json.replace(marker, escape)
                                
                                # Step 5: Restore unicode escapes
                                for marker, unicode_escape in unicode_storage.items():
                                    final_json = final_json.replace(marker, unicode_escape)
                                
                                # Step 6: Try parsing the fixed JSON
                                try:
                                    return json.loads(final_json)
                                except json.JSONDecodeError:
                                    # Still failing, return None to trigger fallback parser
                                    return None
                            else:
                                # Different JSON error, return None to trigger fallback
                                return None
                    
                    # Extract and parse JSON
                    exam_json = extract_and_fix_json(ai_response)
                    
                    if exam_json and isinstance(exam_json, dict) and "questions" in exam_json:
                        st.session_state.generated_exam = exam_json
                        st.session_state.exam_title = exam_title
                        success_msg = t_gen.get('exam_generated', "Examen généré avec succès!")
                        st.success(f"✅ {success_msg}")
                        st.rerun()
                    else:
                        # If JSON parsing failed, try manual extraction
                        st.warning("⚠️ Le format JSON standard n'a pas pu être parsé. Extraction manuelle...")
                        
                        # Try to extract questions manually from the response
                        questions = []
                        
                        # Pattern to find question objects in JSON-like structure
                        # Look for "question": "..." patterns
                        question_pattern = r'"question"\s*:\s*"((?:[^"\\]|\\.|\\n|\\t)*)"'
                        question_matches = re.findall(question_pattern, ai_response, re.DOTALL)
                        
                        # Also look for course patterns
                        course_pattern = r'"course"\s*:\s*"([^"]+)"'
                        course_matches = re.findall(course_pattern, ai_response)
                        
                        if question_matches:
                            # Found questions in JSON format
                            for i, q_text in enumerate(question_matches[:num_questions]):
                                # Unescape the question text
                                q_text = q_text.replace('\\n', '\n').replace('\\t', '\t').replace('\\"', '"')
                                
                                # Get corresponding course
                                course = course_matches[i] if i < len(course_matches) else (courses[i % len(courses)] if courses else "N/A")
                                
                                questions.append({
                                    "course": course,
                                    "question": q_text[:1000],  # Limit length
                                    "answer": None,
                                    "difficulty": difficulty
                                })
                        
                        if questions:
                            exam_data = {
                                "title": exam_title,
                                "subject": subject,
                                "grade": grade,
                                "difficulty": difficulty,
                                "questions": questions
                            }
                            st.session_state.generated_exam = exam_data
                            st.session_state.exam_title = exam_title
                            success_msg = t_gen.get('exam_generated', "Examen généré avec succès!")
                            st.success(f"✅ {success_msg}")
                            st.rerun()
                        else:
                            # Last resort: create basic structure from text
                            st.warning("⚠️ Extraction JSON échouée. Création d'une structure basique...")
                            
                            # Split response into paragraphs as questions
                            paragraphs = [p.strip() for p in ai_response.split('\n\n') if p.strip() and len(p.strip()) > 30]
                            
                            for i, para in enumerate(paragraphs[:num_questions]):
                                # Try to extract question text (skip metadata lines)
                                if not para.startswith(('{', '"', 'title', 'subject', 'grade', 'difficulty', 'questions', 'course')):
                                    questions.append({
                                        "course": courses[i % len(courses)] if courses else "N/A",
                                        "question": para[:1000],
                                        "answer": None,
                                        "difficulty": difficulty
                                    })
                            
                            if questions:
                                exam_data = {
                                    "title": exam_title,
                                    "subject": subject,
                                    "grade": grade,
                                    "difficulty": difficulty,
                                    "questions": questions
                                }
                                st.session_state.generated_exam = exam_data
                                st.session_state.exam_title = exam_title
                                success_msg = t_gen.get('exam_generated', "Examen généré avec succès!")
                                st.success(f"✅ {success_msg}")
                                st.rerun()
                            else:
                                error_msg = t_gen.get('error_extract', "Impossible d'extraire les questions. Veuillez réessayer avec un prompt plus clair.")
                                st.error(f"❌ {error_msg}")
                                raw_msg = t_gen.get('raw_response', "Réponse brute de l'IA:")
                                st.info(raw_msg)
                                st.text_area("", value=ai_response, height=300, key="ai_response_display")
                
                except json.JSONDecodeError as e:
                    error_parse_msg = t_gen.get('error_parse', "Erreur lors du parsing JSON:")
                    st.error(f"❌ {error_parse_msg} {e}")
                    raw_msg = t_gen.get('raw_response', "Réponse brute de l'IA:")
                    st.info(raw_msg)
                    st.text_area("", value=ai_response, height=300, key="ai_response_error")
                except Exception as parse_error:
                    error_extract_msg = t_gen.get('error_extract', "Erreur lors de l'extraction:")
                    st.error(f"❌ {error_extract_msg} {parse_error}")
                    import traceback
                    st.code(traceback.format_exc())
                    raw_msg = t_gen.get('raw_response', "Réponse brute de l'IA:")
                    st.info(raw_msg)
                    st.text_area("", value=ai_response, height=300, key="ai_response_exception")
            else:
                error_detail = res.json().get("detail", res.text) if res.content else res.text
                st.error(f"❌ Erreur lors de la génération: {error_detail}")
        
        except Exception as e:
            st.error(f"❌ Erreur: {str(e)}")
            import traceback
            st.code(traceback.format_exc())



