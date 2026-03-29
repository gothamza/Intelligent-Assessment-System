"""
Exam Taking Page
Allows users to take exams and receive feedback after completion
"""
import streamlit as st
import requests
import os
import json
from datetime import datetime
import re
import io

BACKEND_URL = os.getenv("BACKEND_URL", "http://backend:8000")

def extract_text_from_pdf(uploaded_file):
    """Extract text from PDF file"""
    try:
        import PyPDF2
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text, None
    except ImportError:
        return None, "PyPDF2 n'est pas installé. Veuillez installer: pip install PyPDF2"
    except Exception as e:
        return None, f"Erreur lors de la lecture du PDF: {str(e)}"

def extract_text_from_docx(uploaded_file):
    """Extract text from Word document"""
    try:
        from docx import Document
        doc = Document(uploaded_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text, None
    except ImportError:
        return None, "python-docx n'est pas installé. Veuillez installer: pip install python-docx"
    except Exception as e:
        return None, f"Erreur lors de la lecture du DOCX: {str(e)}"

def parse_questions_from_text(text: str, title: str = "Examen Importé") -> dict:
    """
    Parse questions from plain text using pattern matching.
    Supports formats like:
    - Question 1
      Cours: ... | Difficulté: ...
      Actual question text...
    - Question 1: ...
    - 1. ...
    """
    questions = []
    
    # Try to extract title and duration from document header
    lines = text.strip().split('\n')
    exam_title = title
    exam_duration = None
    
    # Look for title in first few lines
    if lines and (lines[0].strip().lower().startswith(('examen', 'test', 'contrôle', 'évaluation'))):
        exam_title = lines[0].strip()
        text = '\n'.join(lines[1:])
    
    # Try to extract duration from header (common patterns)
    duration_patterns = [
        r'Durée\s*:\s*([^_\n]+)',
        r'Duration\s*:\s*([^_\n]+)',
        r'Temps\s*:\s*([^_\n]+)',
        r'Time\s*:\s*([^_\n]+)',
    ]
    
    for pattern in duration_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            exam_duration = match.group(1).strip()
            exam_duration = re.sub(r'_+', '', exam_duration).strip()
            if exam_duration:
                break
    
    # Pattern for question blocks - match "Question N" followed by "Cours:" metadata
    question_pattern = re.compile(
        r'Question\s+(\d+)\s*\n\s*Cours:\s*([^|\n]+)\s*\|\s*Difficulté:\s*([^\n]+)\s*\n(.+?)(?=\n\s*Question\s+\d+|$)',
        re.MULTILINE | re.DOTALL | re.IGNORECASE
    )
    
    matches = question_pattern.finditer(text)
    
    for match in matches:
        question_num = match.group(1)
        course = match.group(2).strip()
        difficulty = match.group(3).strip()
        block_text = match.group(4).strip()
        
        # Clean the question text
        question_text = block_text
        
        # Remove answer lines (lines with only underscores)
        question_text = re.sub(r'^\s*_{3,}\s*$', '', question_text, flags=re.MULTILINE).strip()
        
        # Remove empty lines
        question_text = '\n'.join([line.strip() for line in question_text.split('\n') if line.strip() and not line.strip().startswith('_')])
        
        # If question is still too short or just underscores, skip it
        if not question_text or len(question_text) < 10 or question_text.replace('_', '').strip() == '':
            continue
        
        questions.append({
            "course": course,
            "difficulty": difficulty,
            "question": question_text
        })
    
    # If no questions found with "Question N" pattern, try numbered format (1., 2., etc.)
    if not questions:
        numbered_pattern = re.compile(
            r'^(\d+)[\.\)]\s*(.+?)(?=^\d+[\.\)]|$)',
            re.MULTILINE | re.DOTALL
        )
        
        numbered_matches = numbered_pattern.finditer(text)
        
        for match in numbered_matches:
            question_text = match.group(2).strip()
            
            # Basic cleaning for numbered questions
            question_text = re.sub(r'^\s*_{3,}\s*$', '', question_text, flags=re.MULTILINE).strip()
            question_text = '\n'.join([line.strip() for line in question_text.split('\n') if line.strip() and not line.strip().startswith('_')])
            
            if not question_text or len(question_text) < 10 or question_text.replace('_', '').strip() == '':
                continue
            
            questions.append({
                "course": "Général",
                "difficulty": "Inconnu",
                "question": question_text
            })

    if not questions:
        return None, "Aucune question n'a pu être extraite du texte."

    return {
        "title": exam_title,
        "subtitle": "Examen Importé",
        "duration": exam_duration,
        "instructions": "Répondez à toutes les questions.",
        "questions": questions
    }, None

def generate_exam_pdf(exam, student_answers, results):
    """Generate PDF from exam results"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, 
                               rightMargin=72, leftMargin=72, 
                               topMargin=72, bottomMargin=72)
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Heading1'],
            fontSize=20,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#1f4788'),
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'HeadingStyle',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.HexColor('#2c3e50'),
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'NormalStyle',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            alignment=TA_LEFT,
            leading=14
        )
        
        question_style = ParagraphStyle(
            'QuestionStyle',
            parent=normal_style,
            fontSize=12,
            fontName='Helvetica-Bold',
            textColor=colors.HexColor('#34495e')
        )
        
        answer_style = ParagraphStyle(
            'AnswerStyle',
            parent=normal_style,
            fontSize=10,
            leftIndent=20,
            textColor=colors.HexColor('#7f8c8d')
        )
        
        feedback_style = ParagraphStyle(
            'FeedbackStyle',
            parent=normal_style,
            fontSize=10,
            leftIndent=20,
            textColor=colors.HexColor('#27ae60'),
            backColor=colors.HexColor('#ecf0f1')
        )
        
        story = []
        
        # Title
        exam_title = exam.get('title', 'Examen')
        story.append(Paragraph(exam_title, title_style))
        story.append(Spacer(1, 10))
        
        # Exam info
        info_data = [
            ['Matière:', exam.get('subject', 'N/A')],
            ['Niveau:', exam.get('grade', 'N/A')],
            ['Difficulté:', exam.get('difficulty', 'N/A')],
            ['Date:', datetime.now().strftime('%d/%m/%Y %H:%M')]
        ]
        info_table = Table(info_data, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#ecf0f1')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 12),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 20))
        
        # Score
        score = float(results.get('score', 0))
        score_color = colors.HexColor('#27ae60') if score >= 70 else (colors.HexColor('#f39c12') if score >= 50 else colors.HexColor('#e74c3c'))
        score_style = ParagraphStyle(
            'ScoreStyle',
            parent=title_style,
            fontSize=18,
            textColor=score_color
        )
        story.append(Paragraph(f"Score Final: {score:.1f}%", score_style))
        story.append(Spacer(1, 20))
        
        # Questions and answers
        questions = exam.get('questions', [])
        classifications = results.get('classifications', [])
        individual_feedbacks = results.get('individual_feedbacks', [])
        
        for i, question_data in enumerate(questions):
            story.append(Paragraph(f"Question {i+1}", heading_style))
            
            # Course and difficulty
            course = question_data.get('course', 'N/A')
            difficulty = question_data.get('difficulty', 'N/A')
            story.append(Paragraph(f"<b>Cours:</b> {course} | <b>Difficulté:</b> {difficulty}", normal_style))
            story.append(Spacer(1, 5))
            
            # Question text
            question_text = question_data.get('question', '').replace('\n', '<br/>')
            story.append(Paragraph(f"<b>Question:</b> {question_text}", question_style))
            story.append(Spacer(1, 5))
            
            # Student answer
            student_answer = student_answers.get(str(i), 'Aucune réponse')
            student_answer_text = str(student_answer).replace('\n', '<br/>')
            story.append(Paragraph(f"<b>Votre réponse:</b>", normal_style))
            story.append(Paragraph(student_answer_text, answer_style))
            story.append(Spacer(1, 5))
            
            # Classification and feedback
            if i < len(classifications):
                classification = classifications[i]
                feedback = individual_feedbacks[i] if i < len(individual_feedbacks) else "Aucun feedback"
                feedback_text = str(feedback).replace('\n', '<br/>')
                
                class_label = classification.replace('_', ' ').title()
                class_color = colors.HexColor('#27ae60') if classification == 'correcte' else (colors.HexColor('#f39c12') if classification == 'partiellement_correcte' else colors.HexColor('#e74c3c'))
                
                story.append(Paragraph(f"<b>Classification:</b> <font color='{class_color.hexval()}'>{class_label}</font>", normal_style))
                story.append(Paragraph(f"<b>Feedback:</b> {feedback_text}", feedback_style))
            
            story.append(Spacer(1, 15))
            
            # Page break every 3 questions
            if (i + 1) % 3 == 0 and i < len(questions) - 1:
                story.append(PageBreak())
        
        # Overall feedback
        if results.get('overall_feedback'):
            story.append(PageBreak())
            story.append(Paragraph("Feedback Global", heading_style))
            overall_feedback = str(results.get('overall_feedback', '')).replace('\n', '<br/>')
            story.append(Paragraph(overall_feedback, normal_style))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        st.error(f"Erreur lors de la génération du PDF: {str(e)}")
        return None

def generate_exam_docx(exam, student_answers, results):
    """Generate DOCX from exam results"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Title
        title = doc.add_heading(exam.get('title', 'Examen'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(20)
        title_run.font.color.rgb = RGBColor(31, 71, 136)
        
        # Exam info
        doc.add_paragraph()
        info_para = doc.add_paragraph()
        info_para.add_run('Matière: ').bold = True
        info_para.add_run(exam.get('subject', 'N/A'))
        info_para.add_run(' | Niveau: ').bold = True
        info_para.add_run(exam.get('grade', 'N/A'))
        info_para.add_run(' | Difficulté: ').bold = True
        info_para.add_run(exam.get('difficulty', 'N/A'))
        info_para.add_run(' | Date: ').bold = True
        info_para.add_run(datetime.now().strftime('%d/%m/%Y %H:%M'))
        
        doc.add_paragraph()
        
        # Score
        score = float(results.get('score', 0))
        score_para = doc.add_paragraph()
        score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        score_run = score_para.add_run(f'Score Final: {score:.1f}%')
        score_run.font.size = Pt(16)
        score_run.bold = True
        if score >= 70:
            score_run.font.color.rgb = RGBColor(39, 174, 96)  # Green
        elif score >= 50:
            score_run.font.color.rgb = RGBColor(243, 156, 18)  # Orange
        else:
            score_run.font.color.rgb = RGBColor(231, 76, 60)  # Red
        
        doc.add_paragraph()
        doc.add_paragraph('─' * 50)
        doc.add_paragraph()
        
        # Questions and answers
        questions = exam.get('questions', [])
        classifications = results.get('classifications', [])
        individual_feedbacks = results.get('individual_feedbacks', [])
        
        for i, question_data in enumerate(questions):
            # Question number
            q_heading = doc.add_heading(f'Question {i+1}', level=2)
            q_heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            
            # Course and difficulty
            course_para = doc.add_paragraph()
            course_para.add_run('Cours: ').bold = True
            course_para.add_run(question_data.get('course', 'N/A'))
            course_para.add_run(' | Difficulté: ').bold = True
            course_para.add_run(question_data.get('difficulty', 'N/A'))
            
            # Question text
            doc.add_paragraph()
            q_para = doc.add_paragraph()
            q_para.add_run('Question: ').bold = True
            q_para.add_run(question_data.get('question', ''))
            
            # Student answer
            doc.add_paragraph()
            ans_para = doc.add_paragraph()
            ans_para.add_run('Votre réponse: ').bold = True
            student_answer = student_answers.get(str(i), 'Aucune réponse')
            ans_para.add_run(str(student_answer))
            ans_para.paragraph_format.left_indent = Inches(0.5)
            
            # Classification and feedback
            if i < len(classifications):
                doc.add_paragraph()
                class_para = doc.add_paragraph()
                class_para.add_run('Classification: ').bold = True
                classification = classifications[i]
                class_label = classification.replace('_', ' ').title()
                class_run = class_para.add_run(class_label)
                
                if classification == 'correcte':
                    class_run.font.color.rgb = RGBColor(39, 174, 96)  # Green
                elif classification == 'partiellement_correcte':
                    class_run.font.color.rgb = RGBColor(243, 156, 18)  # Orange
                else:
                    class_run.font.color.rgb = RGBColor(231, 76, 60)  # Red
                
                # Feedback
                doc.add_paragraph()
                feedback_para = doc.add_paragraph()
                feedback_para.add_run('Feedback: ').bold = True
                feedback = individual_feedbacks[i] if i < len(individual_feedbacks) else "Aucun feedback"
                feedback_para.add_run(str(feedback))
                feedback_para.paragraph_format.left_indent = Inches(0.5)
                feedback_para.paragraph_format.space_before = Pt(6)
                # Add background color (light gray)
                for run in feedback_para.runs:
                    run.font.highlight_color = 7  # Light gray highlight
            
            doc.add_paragraph()
            doc.add_paragraph('─' * 50)
            doc.add_paragraph()
        
        # Overall feedback
        if results.get('overall_feedback'):
            doc.add_page_break()
            overall_heading = doc.add_heading('Feedback Global', level=1)
            overall_heading.runs[0].font.color.rgb = RGBColor(31, 71, 136)
            overall_para = doc.add_paragraph(str(results.get('overall_feedback', '')))
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        st.error(f"Erreur lors de la génération du DOCX: {str(e)}")
        return None

def classify_answer_and_feedback(headers, question, student_answer, language, subject, grade):
    """Classify answer and generate feedback using backend"""
    try:
        # Create prompt for classification and feedback
        prompts = {
            "Français": f"""Évalue cette réponse d'étudiant pour une question de {subject} niveau {grade}.

QUESTION: {question}

RÉPONSE DE L'ÉTUDIANT: {student_answer}

INSTRUCTIONS:
1. Classe la réponse comme: "correcte", "partiellement_correcte", ou "incorrecte"
2. Génère un feedback détaillé et pédagogique

Format de sortie JSON:
{{
  "classification": "correcte|partiellement_correcte|incorrecte",
  "feedback": "Feedback détaillé expliquant pourquoi la réponse est correcte/partiellement correcte/incorrecte, avec suggestions d'amélioration"
}}""",
            
            "English": f"""Evaluate this student answer for a {subject} question at {grade} level.

QUESTION: {question}

STUDENT ANSWER: {student_answer}

INSTRUCTIONS:
1. Classify the answer as: "correct", "partially_correct", or "incorrect"
2. Generate detailed and pedagogical feedback

Output format JSON:
{{
  "classification": "correct|partially_correct|incorrect",
  "feedback": "Detailed feedback explaining why the answer is correct/partially correct/incorrect, with improvement suggestions"
}}""",
            
            "العربية": f"""قيم هذه الإجابة الطلابية لسؤال في {subject} مستوى {grade}.

السؤال: {question}

إجابة الطالب: {student_answer}

التعليمات:
1. صنف الإجابة كـ: "صحيحة"، "صحيحة_جزئيًا"، أو "خاطئة"
2. أنشئ ملاحظات تفصيلية وتعليمية

تنسيق الإخراج JSON:
{{
  "classification": "صحيحة|صحيحة_جزئيًا|خاطئة",
  "feedback": "ملاحظات تفصيلية تشرح لماذا الإجابة صحيحة/صحيحة جزئيًا/خاطئة، مع اقتراحات للتحسين"
}}"""
        }
        
        prompt = prompts.get(language, prompts["Français"])
        
        # Create temporary chat for feedback
        chat_resp = requests.post(
            f"{BACKEND_URL}/chats/create",
            json={"chat_title": f"Feedback Examen {datetime.now().strftime('%Y%m%d_%H%M%S')}"},
            headers=headers
        )
        
        if chat_resp.status_code != 200:
            return "incorrecte", "Erreur lors de la création du chat pour le feedback"
        
        chat_id = chat_resp.json().get("chat_id")
        
        # Call backend for classification and feedback
        body = {
            "doc_ids": [],
            "chat_id": chat_id,
            "messages": [{"role": "user", "content": prompt}],
            "use_web_search": False,
            "model_name": st.session_state.get("selected_model", "Groq/llama-3.1-8b-instant"),
            "subject": subject,
            "grade": grade,
            "course": "",
            "custom_instructions": "",
            "language": language
        }
        
        res = requests.post(
            f"{BACKEND_URL}/rag_langgraph2",
            json=body,
            headers=headers,
            timeout=300
        )
        
        if res.status_code == 200:
            data = res.json()
            ai_response = data.get("answer", "")
            
            # Try to extract JSON from response
            try:
                json_match = re.search(r'\{.*\}', ai_response, re.DOTALL)
                if json_match:
                    result = json.loads(json_match.group())
                    classification = result.get("classification", "incorrecte")
                    feedback = result.get("feedback", "Aucun feedback disponible")
                    
                    # Normalize classification
                    classification_map = {
                        "correcte": "correcte",
                        "correct": "correcte",
                        "partiellement_correcte": "partiellement_correcte",
                        "partiellement": "partiellement_correcte",
                        "partially_correct": "partiellement_correcte",
                        "incorrecte": "incorrecte",
                        "incorrect": "incorrecte",
                        "صحيحة": "correcte",
                        "صحيحة_جزئيًا": "partiellement_correcte",
                        "خاطئة": "incorrecte"
                    }
                    
                    classification = classification_map.get(classification.lower(), "incorrecte")
                    
                    return classification, feedback
                else:
                    # Parse classification from text
                    classification_text = ai_response.lower()
                    if "correct" in classification_text and "incorrect" not in classification_text and "partiellement" not in classification_text:
                        classification = "correcte"
                    elif "partiellement" in classification_text or "partiel" in classification_text:
                        classification = "partiellement_correcte"
                    else:
                        classification = "incorrecte"
                    
                    return classification, ai_response
            except json.JSONDecodeError:
                # Parse from text
                classification_text = ai_response.lower()
                if "correct" in classification_text and "incorrect" not in classification_text and "partiellement" not in classification_text:
                    classification = "correcte"
                elif "partiellement" in classification_text or "partiel" in classification_text:
                    classification = "partiellement_correcte"
                else:
                    classification = "incorrecte"
                
                return classification, ai_response
        else:
            return "incorrecte", f"Erreur backend: {res.text}"
    
    except Exception as e:
        return "incorrecte", f"Erreur: {str(e)}"


def process_exam_results_backend(headers, exam_questions, student_answers):
    """Process exam results using backend"""
    classifications = []
    individual_feedbacks = []
    
    subject = st.session_state.get("selected_subject", "Mathématiques")
    grade = st.session_state.get("selected_grade", "7ème année")
    # Get language dynamically from session state
    language = st.session_state.get("selected_language", "Français")
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for i, question_data in enumerate(exam_questions):
        question = question_data.get("question", "")
        student_answer = student_answers.get(str(i), "")
        
        if student_answer:
            status_text.text(f"Évaluation de la question {i+1}/{len(exam_questions)}...")
            
            # Get language dynamically for each question (in case user changed it)
            current_language = st.session_state.get("selected_language", "Français")
            classification, feedback = classify_answer_and_feedback(
                headers, question, student_answer, current_language, subject, grade
            )
            
            classifications.append(classification)
            individual_feedbacks.append(feedback)
        else:
            classifications.append("incorrecte")
            individual_feedbacks.append("Aucune réponse fournie.")
        
        progress_bar.progress((i + 1) / len(exam_questions))
    
    status_text.empty()
    progress_bar.empty()
    
    # Calculate score
    correct_count = classifications.count("correcte")
    partial_count = classifications.count("partiellement_correcte")
    total_questions = len(classifications)
    
    score = ((correct_count + partial_count * 0.5) / total_questions * 100) if total_questions > 0 else 0
    
    # Generate overall feedback
    try:
        overall_prompt = {
            "Français": f"""Génère un feedback global pour un examen de {subject} niveau {grade}.

RÉSULTATS:
- Score: {score:.1f}%
- Réponses correctes: {correct_count}/{total_questions}
- Réponses partiellement correctes: {partial_count}/{total_questions}
- Réponses incorrectes: {total_questions - correct_count - partial_count}/{total_questions}

Génère un feedback global encourageant et constructif avec des suggestions d'amélioration.""",
            
            "English": f"""Generate overall feedback for a {subject} exam at {grade} level.

RESULTS:
- Score: {score:.1f}%
- Correct answers: {correct_count}/{total_questions}
- Partially correct answers: {partial_count}/{total_questions}
- Incorrect answers: {total_questions - correct_count - partial_count}/{total_questions}

Generate an encouraging and constructive overall feedback with improvement suggestions.""",
            
            "العربية": f"""أنشئ ملاحظات عامة لامتحان في {subject} مستوى {grade}.

النتائج:
- النتيجة: {score:.1f}%
- الإجابات الصحيحة: {correct_count}/{total_questions}
- الإجابات الصحيحة جزئيًا: {partial_count}/{total_questions}
- الإجابات الخاطئة: {total_questions - correct_count - partial_count}/{total_questions}

أنشئ ملاحظات عامة مشجعة وبناءة مع اقتراحات للتحسين."""
        }
        
        # Get language dynamically for overall feedback (in case user changed it)
        current_language = st.session_state.get("selected_language", "Français")
        prompt = overall_prompt.get(current_language, overall_prompt["Français"])
        
        chat_resp = requests.post(
            f"{BACKEND_URL}/chats/create",
            json={"chat_title": f"Feedback Global {datetime.now().strftime('%Y%m%d_%H%M%S')}"},
            headers=headers
        )
        
        if chat_resp.status_code == 200:
            chat_id = chat_resp.json().get("chat_id")
            
            body = {
                "doc_ids": [],
                "chat_id": chat_id,
                "messages": [{"role": "user", "content": prompt}],
                "use_web_search": False,
                "model_name": st.session_state.get("selected_model", "Groq/llama-3.1-8b-instant"),
                "subject": subject,
                "grade": grade,
                "course": "",
                "custom_instructions": "",
                "language": current_language
            }
            
            res = requests.post(
                f"{BACKEND_URL}/rag_langgraph2",
                json=body,
                headers=headers,
                timeout=300
            )
            
            if res.status_code == 200:
                overall_feedback = res.json().get("answer", "Excellent travail! Continue à t'améliorer.")
            else:
                overall_feedback = "Excellent travail! Continue à t'améliorer."
        else:
            overall_feedback = "Excellent travail! Continue à t'améliorer."
    except Exception:
        overall_feedback = "Excellent travail! Continue à t'améliorer."
    
    return classifications, individual_feedbacks, score, overall_feedback


def take_exam_page():
    """Main function for exam taking page"""
    
    st.title("📝 Passer un Examen")
    st.markdown("---")
    st.markdown("Répondez aux questions de l'examen et recevez un feedback détaillé après soumission.")
    
    headers = {"Authorization": f"Bearer {st.session_state.get('token')}"}
    
    # Check authentication
    if not st.session_state.get("token"):
        st.error("❌ Vous devez être connecté pour passer un examen.")
        st.info("Veuillez vous connecter depuis la page d'accueil.")
        return
    
    # Initialize session state
    if "current_exam" not in st.session_state:
        st.session_state.current_exam = None
    if "student_answers" not in st.session_state:
        st.session_state.student_answers = {}
    if "exam_completed" not in st.session_state:
        st.session_state.exam_completed = False
    if "exam_results" not in st.session_state:
        st.session_state.exam_results = None
    if "processed_files" not in st.session_state:
        st.session_state.processed_files = {}
    
    # Clear processed files when a new exam is loaded from session state
    if st.session_state.get("exam_to_take") and not st.session_state.current_exam:
        st.session_state.processed_files = {}
    
    # Exam loading section
    st.subheader("📂 Charger un Examen")
    
    # Option 1: Load from session state (from exam generation page)
    if st.session_state.get("exam_to_take") and not st.session_state.current_exam:
        st.info("✅ Un examen est disponible depuis la page de génération.")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("📋 Charger l'examen disponible", use_container_width=True):
                st.session_state.current_exam = st.session_state.exam_to_take
                st.session_state.exam_to_take = None
                st.session_state.student_answers = {}
                st.session_state.exam_completed = False
                st.session_state.exam_results = None
                st.session_state.processed_files = {}  # Clear processed files
                st.rerun()
        with col2:
            if st.button("🗑️ Ignorer", use_container_width=True):
                st.session_state.exam_to_take = None
                st.rerun()
    
    # Option 2: Upload file (JSON, PDF, or Word)
    uploaded_file = st.file_uploader(
        "Ou téléchargez un fichier d'examen (JSON, PDF, DOCX):",
        type=["json", "pdf", "docx"],
        help="Téléchargez un fichier JSON, PDF ou Word contenant un examen",
        key="exam_file_uploader"
    )
    
    # Process file upload - allow replacing existing exam
    if uploaded_file:
        file_id = f"{uploaded_file.name}_{uploaded_file.size}"
        
        # Check if file already processed (to avoid infinite loop)
        if file_id not in st.session_state.processed_files:
            file_extension = uploaded_file.name.split('.')[-1].lower()
            exam_data = None
            error_msg = None
            
            if file_extension == "json":
                # JSON file
                try:
                    # Reset file pointer
                    uploaded_file.seek(0)
                    exam_data = json.load(uploaded_file)
                    
                    if "questions" in exam_data and isinstance(exam_data["questions"], list):
                        # Ensure all required fields are present with defaults
                        if not exam_data.get("title"):
                            exam_data["title"] = "Examen Importé"
                        if not exam_data.get("subject"):
                            exam_data["subject"] = "N/A"
                        if not exam_data.get("grade"):
                            exam_data["grade"] = "N/A"
                        if not exam_data.get("difficulty"):
                            exam_data["difficulty"] = "N/A"
                        
                        # Clear any existing exam and load new one
                        st.session_state.current_exam = exam_data
                        st.session_state.exam_to_take = None  # Clear exam from generation page
                        st.session_state.student_answers = {}
                        st.session_state.exam_completed = False
                        st.session_state.exam_results = None
                        st.session_state.processed_files[file_id] = True
                        st.success(f"✅ Examen chargé avec succès! {len(exam_data['questions'])} question(s) trouvée(s).")
                        st.rerun()
                    else:
                        st.error("❌ Format d'examen invalide. Le fichier doit contenir une liste 'questions'.")
                        st.session_state.processed_files[file_id] = True  # Mark as processed even if invalid
                except json.JSONDecodeError as e:
                    st.error(f"❌ Erreur lors de la lecture du fichier JSON: {str(e)}")
                    st.session_state.processed_files[file_id] = True
                except Exception as e:
                    st.error(f"❌ Erreur inattendue: {str(e)}")
                    st.session_state.processed_files[file_id] = True
            
            elif file_extension == "pdf":
                # PDF file - extract text then parse
                with st.spinner("📄 Extraction du texte depuis le PDF..."):
                    # Reset file pointer
                    uploaded_file.seek(0)
                    text_content, error_msg = extract_text_from_pdf(uploaded_file)
                    
                    if error_msg:
                        st.error(f"❌ {error_msg}")
                        st.session_state.processed_files[file_id] = True
                    elif text_content:
                        # Show preview
                        with st.expander("📄 Aperçu du texte extrait (premiers 500 caractères)"):
                            st.text(text_content[:500])
                        
                        # Parse questions from text
                        with st.spinner("🔍 Analyse des questions..."):
                            exam_data, parse_error = parse_questions_from_text(text_content, uploaded_file.name)
                            
                            if parse_error:
                                st.error(f"❌ {parse_error}")
                                st.session_state.processed_files[file_id] = True
                            elif exam_data and exam_data.get("questions"):
                                st.session_state.current_exam = exam_data
                                st.session_state.student_answers = {}
                                st.session_state.exam_completed = False
                                st.session_state.exam_results = None
                                st.session_state.processed_files[file_id] = True
                                st.success(f"✅ {len(exam_data['questions'])} question(s) extraite(s) du PDF!")
                                st.rerun()
                            else:
                                st.error("❌ Aucune question trouvée dans le PDF. Vérifiez le format du document.")
                                st.info("💡 Le format attendu est: 'Question N\nCours: ... | Difficulté: ...\nTexte de la question...'")
                                st.session_state.processed_files[file_id] = True
            
            elif file_extension in ["docx", "doc"]:
                # Word file - extract text then parse
                with st.spinner("📘 Extraction du texte depuis le document Word..."):
                    # Reset file pointer
                    uploaded_file.seek(0)
                    text_content, error_msg = extract_text_from_docx(uploaded_file)
                    
                    if error_msg:
                        st.error(f"❌ {error_msg}")
                        st.session_state.processed_files[file_id] = True
                    elif text_content:
                        # Show preview
                        with st.expander("📘 Aperçu du texte extrait (premiers 500 caractères)"):
                            st.text(text_content[:500])
                        
                        # Parse questions from text
                        with st.spinner("🔍 Analyse des questions..."):
                            exam_data, parse_error = parse_questions_from_text(text_content, uploaded_file.name)
                            
                            if parse_error:
                                st.error(f"❌ {parse_error}")
                                st.session_state.processed_files[file_id] = True
                            elif exam_data and exam_data.get("questions"):
                                st.session_state.current_exam = exam_data
                                st.session_state.student_answers = {}
                                st.session_state.exam_completed = False
                                st.session_state.exam_results = None
                                st.session_state.processed_files[file_id] = True
                                st.success(f"✅ {len(exam_data['questions'])} question(s) extraite(s) du document Word!")
                                st.rerun()
                            else:
                                st.error("❌ Aucune question trouvée dans le document Word. Vérifiez le format du document.")
                                st.info("💡 Le format attendu est: 'Question N\nCours: ... | Difficulté: ...\nTexte de la question...'")
                                st.session_state.processed_files[file_id] = True
            
            else:
                st.error(f"❌ Type de fichier non supporté: {file_extension}")
                st.session_state.processed_files[file_id] = True
    
    # Display exam and collect answers
    exam = st.session_state.current_exam
    
    if exam and not st.session_state.exam_completed:
        st.markdown("---")
        st.subheader("📋 Questions de l'Examen")
        
        # Exam header
        st.markdown(f"### {exam.get('title', 'Examen')}")
        st.markdown(f"**Matière:** {exam.get('subject', 'N/A')} | **Niveau:** {exam.get('grade', 'N/A')} | **Difficulté:** {exam.get('difficulty', 'N/A')}")
        st.markdown(f"**Nombre de questions:** {len(exam.get('questions', []))}")
        st.markdown("---")
        
        # Display questions and collect answers
        questions = exam.get('questions', [])
        
        for i, question_data in enumerate(questions):
            course = question_data.get('course', 'N/A')
            question_text = question_data.get('question', '')
            difficulty = question_data.get('difficulty', 'N/A')
            
            st.markdown(f"#### Question {i+1}")
            st.markdown(f"**Cours:** {course} | **Difficulté:** {difficulty}")
            st.markdown(f"**{question_text}**")
            
            # Answer input
            answer_key = f"answer_{i}"
            if answer_key not in st.session_state.student_answers:
                st.session_state.student_answers[str(i)] = ""
            
            st.session_state.student_answers[str(i)] = st.text_area(
                f"Votre réponse pour la question {i+1}:",
                value=st.session_state.student_answers.get(str(i), ""),
                key=answer_key,
                height=100,
                help="Écrivez votre réponse à cette question"
            )
            
            st.markdown("---")
        
        # Submit button
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button("✅ Soumettre l'Examen", type="primary", use_container_width=True):
                # Check if all questions are answered
                unanswered = [i for i in range(len(questions)) if not st.session_state.student_answers.get(str(i), "").strip()]
                
                if unanswered:
                    st.warning(f"⚠️ Veuillez répondre à toutes les questions. Questions non répondues: {[q+1 for q in unanswered]}")
                else:
                    # Process exam results
                    with st.spinner("🤖 Évaluation de l'examen en cours..."):
                        classifications, individual_feedbacks, score, overall_feedback = process_exam_results_backend(
                            headers, questions, st.session_state.student_answers
                        )
                        
                        st.session_state.exam_results = {
                            "classifications": classifications,
                            "individual_feedbacks": individual_feedbacks,
                            "score": score,
                            "overall_feedback": overall_feedback
                        }
                        
                        st.session_state.exam_completed = True
                        st.rerun()
    
    # Display results after exam completion
    if st.session_state.exam_completed and st.session_state.exam_results:
        st.markdown("---")
        st.header("🎉 Examen Terminé!")
        
        results = st.session_state.exam_results
        score = float(results.get("score", 0))
        
        # Score display with color
        color = "green" if score >= 70 else ("orange" if score >= 50 else "red")
        emoji = "🎉" if score >= 70 else ("👍" if score >= 50 else "💪")
        
        st.markdown(f"""
        <div style="text-align:center;padding:30px;background:linear-gradient(135deg, #667eea 0%, #764ba2 100%);border-radius:15px;margin:20px 0;color:white;">
            <h1 style="color:white;margin:0;">{emoji} Score Final: {score:.1f}%</h1>
            <p style="font-size:18px;margin-top:10px;">Excellent travail! Continue à t'améliorer.</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Individual question feedback
        if results.get("classifications") and exam:
            st.markdown("---")
            st.subheader("📝 Feedback par Question")
            
            questions = exam.get('questions', [])
            classifications = results.get("classifications", [])
            individual_feedbacks = results.get("individual_feedbacks", [])
            
            for i, question_data in enumerate(questions):
                if i < len(classifications):
                    classification = classifications[i]
                    feedback = individual_feedbacks[i] if i < len(individual_feedbacks) else "Aucun feedback disponible"
                    student_answer = st.session_state.student_answers.get(str(i), "Aucune réponse")
                    
                    # Classification color
                    if classification == "correcte":
                        class_color = "green"
                        class_icon = "✅"
                    elif classification == "partiellement_correcte":
                        class_color = "orange"
                        class_icon = "⚠️"
                    else:
                        class_color = "red"
                        class_icon = "❌"
                    
                    with st.expander(f"{class_icon} Question {i+1} - {classification.replace('_', ' ').title()}", expanded=(i == 0)):
                        st.markdown(f"**Question:** {question_data.get('question', 'N/A')}")
                        st.markdown(f"**Votre réponse:** {student_answer}")
                        st.markdown(f"**Classification:** :{class_color}[{classification.replace('_', ' ').title()}]")
                        st.markdown(f"**Feedback:**\n{feedback}")
        
        # Overall feedback
        if results.get("overall_feedback"):
            st.markdown("---")
            st.subheader("💬 Feedback Global")
            st.info(results.get("overall_feedback"))
        
        # Download results
        st.markdown("---")
        st.subheader("💾 Télécharger les Résultats")
        
        col_d1, col_d2, col_d3, col_d4 = st.columns(4)
        
        with col_d1:
            # Download results as JSON
            results_json = {
                "exam": exam,
                "student_answers": st.session_state.student_answers,
                "results": results,
                "date_completed": datetime.now().isoformat()
            }
            
            st.download_button(
                label="📄 JSON",
                data=json.dumps(results_json, ensure_ascii=False, indent=2),
                file_name=f"resultats_examen_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json",
                use_container_width=True,
                help="Télécharger les résultats au format JSON"
            )
        
        with col_d2:
            # Download as PDF
            pdf_data = generate_exam_pdf(exam, st.session_state.student_answers, results)
            if pdf_data:
                st.download_button(
                    label="📕 PDF",
                    data=pdf_data,
                    file_name=f"resultats_examen_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    help="Télécharger les résultats au format PDF"
                )
            else:
                st.button("📕 PDF", disabled=True, use_container_width=True, help="Erreur lors de la génération")
        
        with col_d3:
            # Download as DOCX
            docx_data = generate_exam_docx(exam, st.session_state.student_answers, results)
            if docx_data:
                st.download_button(
                    label="📘 Word",
                    data=docx_data,
                    file_name=f"resultats_examen_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    help="Télécharger les résultats au format Word (DOCX)"
                )
            else:
                st.button("📘 Word", disabled=True, use_container_width=True, help="Erreur lors de la génération")
        
        with col_d4:
            # Reset button to take another exam
            if st.button("🔄 Nouvel Examen", use_container_width=True):
                st.session_state.current_exam = None
                st.session_state.student_answers = {}
                st.session_state.exam_completed = False
                st.session_state.exam_results = None
                st.session_state.exam_to_take = None
                st.rerun()



