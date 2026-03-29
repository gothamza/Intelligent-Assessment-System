"""
Export helpers for exams (DOCX/PDF/Text/JSON).
"""
from io import BytesIO


def export_docx_from_html(html_content: str) -> BytesIO | None:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re

        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Extract title
        title_match = re.search(r'<div class="exam-title">([^<]+)</div>', html_content)
        if title_match:
            p = doc.add_heading(title_match.group(1), 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        
        # Extract exam info
        info_match = re.search(r'<div class="exam-info">([^<]+)</div>', html_content)
        if info_match:
            p = doc.add_paragraph(info_match.group(1))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].font.size = Pt(12)
        
        # Add student info section
        doc.add_paragraph()
        for sl in [
            "Nom: _________________________________",
            "Prénom: _______________________________",
            "Classe: _______________________________",
        ]:
            doc.add_paragraph(sl)
        doc.add_paragraph()
        
        # Extract all questions using a more robust approach
        # Split by question divs and process each one
        question_sections = re.split(r'<div class="question">', html_content)
        
        for i, section in enumerate(question_sections[1:], 1):  # Skip first empty section
            # Find the end of this question section
            end_match = re.search(r'</div>\s*</div>', section)
            if end_match:
                question_html = section[:end_match.start()]
            else:
                question_html = section
            
            # Extract question number
            number_match = re.search(r'<div class="question-number">(.*?)</div>', question_html)
            if number_match:
                p = doc.add_heading(re.sub(r'<[^>]+>', '', number_match.group(1)).strip(), level=2)
                if p.runs:
                    p.runs[0].font.size = Pt(14)
                    p.runs[0].bold = True
            
            # Extract question meta
            meta_match = re.search(r'<div class="question-meta">(.*?)</div>', question_html)
            if meta_match:
                p = doc.add_paragraph(re.sub(r'<[^>]+>', '', meta_match.group(1)).strip())
                if p.runs:
                    p.runs[0].italic = True
                    p.runs[0].font.size = Pt(10)
            
            # Extract question text - this is the most important part
            text_match = re.search(r'<div class="question-text">(.*?)</div>', question_html, re.DOTALL)
            if text_match:
                text = re.sub(r'<[^>]+>', '', text_match.group(1)).strip()
                if text:  # Only add if text exists
                    doc.add_paragraph(text)
            else:
                # Fallback: try to find any text content in the question
                fallback_text = re.sub(r'<[^>]+>', '', question_html).strip()
                if fallback_text and len(fallback_text) > 10:
                    doc.add_paragraph(fallback_text)
            
            # Add answer space
            doc.add_paragraph()
            for _ in range(3):
                p = doc.add_paragraph("_" * 60)
                if p.runs:
                    p.runs[0].font.size = Pt(12)
            doc.add_paragraph()
            
            # Extract answer if present - improved pattern matching
            answer_match = re.search(r'<div class="answer">(.*?)</div>\s*</div>', question_html, re.DOTALL)
            if not answer_match:
                # Try alternative pattern
                answer_match = re.search(r'<div class="answer">(.*?)(?=</div>\s*</div>|$)', question_html, re.DOTALL)
            
            if answer_match:
                answer_html = answer_match.group(1)
                label_match = re.search(r'<div class="answer-label">([^<]+)</div>', answer_html)
                if label_match:
                    doc.add_heading(label_match.group(1), level=3)
                
                # Extract answer text (everything after the label)
                answer_text = re.sub(r'<div class="answer-label">[^<]+</div>', '', answer_html)
                answer_text = re.sub(r'<[^>]+>', '', answer_text).strip()
                if answer_text:
                    doc.add_paragraph(answer_text)
                else:
                    # Fallback: if no label found, use all the answer content
                    fallback_answer = re.sub(r'<[^>]+>', '', answer_html).strip()
                    if fallback_answer:
                        doc.add_heading("Réponse:", level=3)
                        doc.add_paragraph(fallback_answer)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
    except Exception:
        return None


def export_pdf_from_html(html_content: str) -> BytesIO | None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        import re

        buf = BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm, leftMargin=2*cm, rightMargin=2*cm)
        styles = getSampleStyleSheet()
        story = []

        title_style = ParagraphStyle('ExamTitle', parent=styles['Heading1'], fontSize=18, spaceAfter=20, alignment=TA_CENTER)

        # Extract title
        title_match = re.search(r'<div class="exam-title">([^<]+)</div>', html_content)
        if title_match:
            story.append(Paragraph(title_match.group(1), title_style))
            story.append(Spacer(1, 20))
        
        # Extract exam info
        info_match = re.search(r'<div class="exam-info">([^<]+)</div>', html_content)
        if info_match:
            story.append(Paragraph(info_match.group(1), styles['Normal']))
        
        # Add student info section
        story.append(Spacer(1, 20))
        tbl = Table([["Nom: _________________________________"], ["Prénom: _______________________________"], ["Classe: _______________________________"]], colWidths=[15*cm])
        tbl.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, colors.grey)]))
        story.append(tbl)
        story.append(Spacer(1, 20))
        
        # Extract all questions using a more robust approach
        # Split by question divs and process each one
        question_sections = re.split(r'<div class="question">', html_content)
        
        for i, section in enumerate(question_sections[1:], 1):  # Skip first empty section
            # Find the end of this question section
            end_match = re.search(r'</div>\s*</div>', section)
            if end_match:
                question_html = section[:end_match.start()]
            else:
                question_html = section
            
            # Extract question number
            number_match = re.search(r'<div class="question-number">(.*?)</div>', question_html)
            if number_match:
                story.append(Paragraph(re.sub(r'<[^>]+>', '', number_match.group(1)).strip(), styles['Heading2']))
            
            # Extract question meta
            meta_match = re.search(r'<div class="question-meta">(.*?)</div>', question_html)
            if meta_match:
                story.append(Paragraph(re.sub(r'<[^>]+>', '', meta_match.group(1)).strip(), styles['Normal']))
            
            # Extract question text - this is the most important part
            text_match = re.search(r'<div class="question-text">(.*?)</div>', question_html, re.DOTALL)
            if text_match:
                text = re.sub(r'<[^>]+>', '', text_match.group(1)).strip()
                if text:  # Only add if text exists
                    story.append(Paragraph(text, styles['Normal']))
                    story.append(Spacer(1, 10))
            else:
                # Fallback: try to find any text content in the question
                fallback_text = re.sub(r'<[^>]+>', '', question_html).strip()
                if fallback_text and len(fallback_text) > 10:
                    story.append(Paragraph(fallback_text, styles['Normal']))
                    story.append(Spacer(1, 10))
            
            # Add answer space
            story.append(Paragraph("_" * 80, styles['Normal']))
            story.append(Spacer(1, 10))
            
            # Extract answer if present - improved pattern matching
            answer_match = re.search(r'<div class="answer">(.*?)</div>\s*</div>', question_html, re.DOTALL)
            if not answer_match:
                # Try alternative pattern
                answer_match = re.search(r'<div class="answer">(.*?)(?=</div>\s*</div>|$)', question_html, re.DOTALL)
            
            if answer_match:
                answer_html = answer_match.group(1)
                label_match = re.search(r'<div class="answer-label">([^<]+)</div>', answer_html)
                if label_match:
                    story.append(Paragraph(label_match.group(1), styles['Heading3']))
                
                # Extract answer text (everything after the label)
                answer_text = re.sub(r'<div class="answer-label">[^<]+</div>', '', answer_html)
                answer_text = re.sub(r'<[^>]+>', '', answer_text).strip()
                if answer_text:
                    story.append(Paragraph(answer_text, styles['Normal']))
                else:
                    # Fallback: if no label found, use all the answer content
                    fallback_answer = re.sub(r'<[^>]+>', '', answer_html).strip()
                    if fallback_answer:
                        story.append(Paragraph("Réponse:", styles['Heading3']))
                        story.append(Paragraph(fallback_answer, styles['Normal']))

        doc.build(story)
        buf.seek(0)
        return buf
    except Exception:
        return None


def export_course_to_docx(course_data):
    """Export personalized course to DOCX format"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        
        doc = Document()
        
        # Title
        title = doc.add_heading(course_data.get('title', 'Cours Personnalisé'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Introduction
        doc.add_heading('Introduction', level=1)
        doc.add_paragraph(course_data.get('introduction', ''))
        
        # Topics
        for topic in course_data.get('topics', []):
            doc.add_heading(topic.get('name', 'Sujet'), level=1)
            doc.add_paragraph(topic.get('explanation', ''))
            
            # Exercises
            doc.add_heading('Exercices Pratiques', level=2)
            for i, exercise in enumerate(topic.get('exercises', []), 1):
                doc.add_heading(f'Exercice {i}', level=3)
                doc.add_paragraph(f"Énoncé: {exercise.get('question', '')}")
                doc.add_paragraph(f"Solution: {exercise.get('solution', '')}")
                
                if exercise.get('hints'):
                    doc.add_paragraph("Indices:")
                    for hint in exercise['hints']:
                        doc.add_paragraph(f"• {hint}", style='List Bullet')
        
        # Tips
        if course_data.get('tips'):
            doc.add_heading('Conseils d\'apprentissage', level=1)
            for tip in course_data['tips']:
                doc.add_paragraph(f"• {tip}", style='List Bullet')
        
        # Resources
        if course_data.get('resources'):
            doc.add_heading('Ressources supplémentaires', level=1)
            for resource in course_data['resources']:
                doc.add_paragraph(f"• {resource}", style='List Bullet')
        
        # Save to buffer
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
    except Exception as e:
        raise Exception(f"Erreur lors de la génération du DOCX: {str(e)}")


def export_course_to_pdf(course_data):
    """Export personalized course to PDF format"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        import io
        
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=72, leftMargin=72, 
                               topMargin=72, bottomMargin=18)
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1,  # Center alignment
            textColor=colors.darkblue
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            textColor=colors.darkblue
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            leading=14
        )
        
        story = []
        
        # Title
        story.append(Paragraph(course_data.get('title', 'Cours Personnalisé'), title_style))
        story.append(Spacer(1, 20))
        
        # Introduction
        story.append(Paragraph('Introduction', heading_style))
        story.append(Paragraph(course_data.get('introduction', ''), normal_style))
        story.append(Spacer(1, 20))
        
        # Topics
        for topic in course_data.get('topics', []):
            story.append(Paragraph(topic.get('name', 'Sujet'), heading_style))
            story.append(Paragraph(topic.get('explanation', ''), normal_style))
            story.append(Spacer(1, 12))
            
            # Exercises
            story.append(Paragraph('Exercices Pratiques', heading_style))
            for i, exercise in enumerate(topic.get('exercises', []), 1):
                story.append(Paragraph(f'Exercice {i}', styles['Heading3']))
                
                story.append(Paragraph(f"<b>Énoncé:</b> {exercise.get('question', '')}", normal_style))
                story.append(Paragraph(f"<b>Solution:</b> {exercise.get('solution', '')}", normal_style))
                
                if exercise.get('hints'):
                    story.append(Paragraph("<b>Indices:</b>", normal_style))
                    for hint in exercise['hints']:
                        story.append(Paragraph(f"• {hint}", normal_style))
                
                story.append(Spacer(1, 12))
        
        # Tips
        if course_data.get('tips'):
            story.append(Paragraph('Conseils d\'apprentissage', heading_style))
            for tip in course_data['tips']:
                story.append(Paragraph(f"• {tip}", normal_style))
            story.append(Spacer(1, 20))
        
        # Resources
        if course_data.get('resources'):
            story.append(Paragraph('Ressources supplémentaires', heading_style))
            for resource in course_data['resources']:
                story.append(Paragraph(f"• {resource}", normal_style))
        
        doc.build(story)
        buf.seek(0)
        return buf.getvalue()
    except Exception as e:
        raise Exception(f"Erreur lors de la génération du PDF: {str(e)}")


